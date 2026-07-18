"""تصدير العرض إلى ملف Word احترافي بهوية عزوم — عربي RTL كامل."""
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

from .config import BRAND

PRIMARY = RGBColor.from_string(BRAND["primary"])
ACCENT = RGBColor.from_string(BRAND["accent"])
FONT = "Sakkal Majalla"
FONT_FALLBACK = "Arial"


def _set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _run(paragraph, text, size=13, bold=False, color=None):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_FALLBACK
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), FONT)
    rFonts.set(qn("w:ascii"), FONT_FALLBACK)
    cs_bold = OxmlElement("w:bCs")
    cs_bold.set(qn("w:val"), "1" if bold else "0")
    rPr.append(cs_bold)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(size * 2))
    rPr.append(sz_cs)
    if color:
        run.font.color.rgb = color
    return run


def _para(doc, text="", size=13, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    _set_rtl(p)
    if text:
        _run(p, text, size=size, bold=bold, color=color)
    return p


def _heading(doc, text, level=1):
    if level == 1:
        p = _para(doc, text, size=17, bold=True, color=PRIMARY, space_after=10)
        _add_bottom_border(p)
    else:
        _para(doc, text, size=14, bold=True, color=ACCENT, space_after=8)


def _add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), BRAND["accent"])
    borders.append(bottom)
    pPr.append(borders)


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _table(doc, headers, rows, widths=None, money_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # اتجاه الجدول من اليمين لليسار
    tblPr = table._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_rtl(p)
        _run(p, h, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(cell, BRAND["primary"])

    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_rtl(p)
            txt = f"{value:,.2f}" if i in money_cols and isinstance(value, (int, float)) else str(value)
            _run(p, txt, size=11)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def _cover_page(doc, proposal, settings):
    for _ in range(4):
        doc.add_paragraph()
    _para(doc, BRAND["name_ar"], size=36, bold=True, color=PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, BRAND["name_en"], size=16, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    _para(doc, "العرض الفني والمالي", size=26, bold=True, color=PRIMARY, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, proposal["title"], size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    _para(doc, f"مقدم إلى: {proposal['client']}", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f"رقم العرض: {proposal['ref_no']}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f"تاريخ التقديم: {proposal['created_at'][:10]}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f"سريان العرض: {settings.get('validity_days', '90')} يوماً", size=12,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    _para(doc, settings.get("company_name", ""), size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    contact = " — ".join(x for x in (settings.get("company_phone"), settings.get("company_email")) if x)
    if contact:
        _para(doc, contact, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def export_proposal_docx(proposal: dict, settings: dict, path: str):
    data = proposal["data"]
    doc = Document()

    # هوامش وإعداد عام
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.right_margin = Cm(2.2)
        section.left_margin = Cm(2.2)

    _cover_page(doc, proposal, settings)

    # ---------- الجزء الأول: العرض الفني ----------
    _heading(doc, "أولاً: العرض الفني")
    for sec in data.get("technical_sections", []):
        _heading(doc, sec["title"], level=2)
        for paragraph_text in sec["body"].split("\n"):
            if paragraph_text.strip():
                _para(doc, paragraph_text.strip(), size=12)

    scope = data.get("scope") or []
    if scope and len(scope) > 1:
        _heading(doc, "نطاق العمل التفصيلي", level=2)
        for item in scope:
            _para(doc, f"•  {item}", size=12)

    team = data.get("team") or []
    if team:
        _heading(doc, "فريق العمل المقترح", level=2)
        _table(doc, ["الدور الوظيفي", "العدد"],
               [(t["role"], int(t["count"])) for t in team], widths=[10, 4])

    matrix = data.get("compliance_matrix") or []
    if matrix:
        _heading(doc, "مصفوفة الالتزام بالمتطلبات", level=2)
        _table(doc, ["المتطلب", "الالتزام", "الموضع في العرض"],
               [(m["requirement"], m["response"], m["reference"]) for m in matrix],
               widths=[8, 3, 5])

    doc.add_page_break()

    # ---------- الجزء الثاني: الخطة التنفيذية ----------
    _heading(doc, "ثانياً: الخطة التنفيذية للمشروع")
    _para(doc, f"المدة الإجمالية المقترحة: {int(data.get('duration_weeks', 0))} أسبوعاً", size=13, bold=True)
    for i, phase in enumerate(data.get("plan", []), 1):
        _heading(doc, f"المرحلة {i}: {phase['phase']} ({int(phase['duration_weeks'])} أسابيع)", level=2)
        _para(doc, phase["description"], size=12)
        for d in phase.get("deliverables", []):
            _para(doc, f"–  {d}", size=11)

    doc.add_page_break()

    # ---------- الجزء الثالث: العرض المالي ----------
    _heading(doc, "ثالثاً: العرض المالي")
    _heading(doc, "جدول الكميات والأسعار", level=2)
    boq = data.get("boq", [])
    _table(
        doc,
        ["م", "الكود", "البند", "الوحدة", "الكمية", "سعر الوحدة (ر.س)", "الإجمالي (ر.س)"],
        [(i, l.get("code", ""), l["name"], l["unit"], l["qty"], l["unit_price"], l["total"])
         for i, l in enumerate(boq, 1)],
        widths=[1, 2, 6, 2, 2, 3, 3],
        money_cols=(5, 6),
    )

    fin = data.get("financial", {})
    _heading(doc, "ملخص القيمة الإجمالية", level=2)
    _table(
        doc,
        ["البيان", "القيمة (ريال سعودي)"],
        [
            ("التكلفة المباشرة (جدول الكميات)", fin.get("direct_cost", 0)),
            (f"المصاريف الإدارية والعمومية ({fin.get('overhead_pct', 0):g}%)", fin.get("overhead", 0)),
            (f"احتياطي المخاطر ({fin.get('risk_pct', 0):g}%)", fin.get("risk", 0)),
            (f"هامش الربح ({fin.get('profit_pct', 0):g}%)", fin.get("profit", 0)),
            ("الإجمالي قبل الضريبة", fin.get("subtotal", 0)),
            (f"ضريبة القيمة المضافة ({fin.get('vat_rate', 15):g}%)", fin.get("vat", 0)),
            ("الإجمالي النهائي شامل الضريبة", fin.get("grand_total", 0)),
        ],
        widths=[9, 5],
        money_cols=(1,),
    )

    _heading(doc, "شروط الدفع", level=2)
    _para(doc, settings.get("payment_terms", ""), size=12)

    assumptions = data.get("assumptions") or []
    if assumptions:
        _heading(doc, "الافتراضات والاستثناءات", level=2)
        for a in assumptions:
            _para(doc, f"•  {a}", size=11)

    _para(doc)
    _para(doc, "وتفضلوا بقبول فائق الاحترام والتقدير،", size=12)
    _para(doc, settings.get("company_name", "شركة عزوم"), size=13, bold=True, color=PRIMARY)

    doc.save(path)
    return path
