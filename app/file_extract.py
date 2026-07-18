"""استخراج النصوص من ملفات المشروع المرفوعة (PDF / DOCX / XLSX / TXT)."""
from io import BytesIO
from pathlib import Path

MAX_CHARS_PER_FILE = 60_000


def extract_text(filename: str, content: bytes) -> str:
    ext = Path(filename).suffix.lower()
    try:
        if ext == ".pdf":
            text = _extract_pdf(content)
        elif ext == ".docx":
            text = _extract_docx(content)
        elif ext in (".xlsx", ".xlsm"):
            text = _extract_xlsx(content)
        elif ext in (".txt", ".md", ".csv"):
            text = content.decode("utf-8", errors="replace")
        else:
            return f"[تنسيق غير مدعوم: {ext}]"
    except Exception as exc:  # ملف تالف أو محمي — نُبلغ بدل أن نفشل
        return f"[تعذر استخراج النص من {filename}: {exc}]"
    text = text.strip()
    if len(text) > MAX_CHARS_PER_FILE:
        text = text[:MAX_CHARS_PER_FILE] + "\n...[تم اقتطاع بقية الملف]"
    return text


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx(content: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"## ورقة: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
